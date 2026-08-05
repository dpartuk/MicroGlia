import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

def create_research_proposal():
    doc = docx.Document()

    # Page Margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Styles & Colors
    NAVY = RGBColor(0, 51, 102)     # #003366
    SLATE = RGBColor(70, 80, 95)    # #46505F
    DARK = RGBColor(30, 30, 30)     # #1E1E1E

    # Helper Functions
    def set_cell_background(cell, fill_hex):
        tcPr = cell._tc.get_or_add_tcPr()
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
        tcPr.append(shd)

    def set_cell_margins(cell, top=30, bottom=30, left=80, right=80):
        tcPr = cell._tc.get_or_add_tcPr()
        tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
        tcPr.append(tcMar)

    def add_title(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(22)
        run.font.bold = True
        run.font.color.rgb = NAVY
        p.paragraph_format.space_before = Pt(24)
        p.paragraph_format.space_after = Pt(12)
        return p

    def add_subtitle(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(14)
        run.font.italic = True
        run.font.color.rgb = SLATE
        p.paragraph_format.space_after = Pt(24)
        return p

    def add_h1(text):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.color.rgb = NAVY
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.keep_with_next = True
        return p

    def add_h2(text):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(13)
        run.font.bold = True
        run.font.color.rgb = SLATE
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True
        return p

    def add_body(text, bold_prefix=""):
        p = doc.add_paragraph()
        if bold_prefix:
            r_bold = p.add_run(bold_prefix)
            r_bold.font.name = 'Arial'
            r_bold.font.size = Pt(11)
            r_bold.font.bold = True
            r_bold.font.color.rgb = DARK
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(11)
        run.font.color.rgb = DARK
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.space_after = Pt(6)
        return p

    def add_bullet(text, bold_prefix=""):
        p = doc.add_paragraph(style='List Bullet')
        if bold_prefix:
            r_bold = p.add_run(bold_prefix)
            r_bold.font.name = 'Arial'
            r_bold.font.size = Pt(11)
            r_bold.font.bold = True
            r_bold.font.color.rgb = DARK
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(11)
        run.font.color.rgb = DARK
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.space_after = Pt(4)
        return p

    def add_image_figure(image_path, caption_text, width_inches=6.0):
        if os.path.exists(image_path):
            p_img = doc.add_paragraph()
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_img.paragraph_format.space_before = Pt(8)
            p_img.paragraph_format.space_after = Pt(4)
            run = p_img.add_run()
            run.add_picture(image_path, width=Inches(width_inches))
            
            p_cap = doc.add_paragraph()
            p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_cap.paragraph_format.space_after = Pt(12)
            r_cap = p_cap.add_run(caption_text)
            r_cap.font.name = 'Arial'
            r_cap.font.size = Pt(9.5)
            r_cap.font.italic = True
            r_cap.font.color.rgb = SLATE

    # TITLE PAGE / METADATA
    p_inst = doc.add_paragraph()
    p_inst.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_inst = p_inst.add_run("SCHOOL OF DATA SCIENCE: INTELLIGENT SYSTEMS\nAFEKA ACADEMIC COLLEGE OF ENGINEERING / HEBREW UNIVERSITY OF JERUSALEM")
    r_inst.font.name = 'Arial'
    r_inst.font.size = Pt(11)
    r_inst.font.bold = True
    r_inst.font.color.rgb = SLATE
    p_inst.paragraph_format.space_before = Pt(36)
    p_inst.paragraph_format.space_after = Pt(24)

    add_title("Topological AI Pipeline for Microglial Morphological Classification and Activation Scoring")
    add_subtitle("A Research Proposal Submitted toward the Degree of Master of Science (M.Sc.) in Intelligent Systems")

    # Meta Table
    meta_table = doc.add_table(rows=4, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_data = [
        ("Student Name:", "Doron Peleg"),
        ("Academic Program:", "M.Sc. in Intelligent Systems (Machine Learning)"),
        ("Supervisors:", "Dr. Hadas Lapid\nDr. Lilach Gavish, PhD, MPH\nReut Zinger"),
        ("Submission Date:", "August 2026")
    ]
    for row_idx, (label, val) in enumerate(meta_data):
        row = meta_table.rows[row_idx]
        cell_lbl, cell_val = row.cells[0], row.cells[1]
        
        r_l = cell_lbl.paragraphs[0].add_run(label)
        r_l.font.name = 'Arial'
        r_l.font.bold = True
        r_l.font.size = Pt(11)
        r_l.font.color.rgb = NAVY
        
        r_v = cell_val.paragraphs[0].add_run(val)
        r_v.font.name = 'Arial'
        r_v.font.size = Pt(11)
        r_v.font.color.rgb = DARK
        
        set_cell_margins(cell_lbl, top=60, bottom=60, left=100, right=100)
        set_cell_margins(cell_val, top=60, bottom=60, left=100, right=100)

    doc.add_page_break()

    # TABLE OF ABBREVIATIONS (COMPACT SINGLE-PAGE FIT)
    h_abbr = add_h1("Table of Abbreviations")
    h_abbr.paragraph_format.space_before = Pt(0)
    h_abbr.paragraph_format.space_after = Pt(4)

    abbrev_table = doc.add_table(rows=1, cols=2)
    abbrev_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    hdr_cells = abbrev_table.rows[0].cells
    p_h0 = hdr_cells[0].paragraphs[0]
    p_h1 = hdr_cells[1].paragraphs[0]
    p_h0.paragraph_format.space_before = Pt(2)
    p_h0.paragraph_format.space_after = Pt(2)
    p_h1.paragraph_format.space_before = Pt(2)
    p_h1.paragraph_format.space_after = Pt(2)

    r_h0 = p_h0.add_run("Abbreviation")
    r_h1 = p_h1.add_run("Definition")
    r_h0.font.name = 'Arial'
    r_h0.font.bold = True
    r_h0.font.size = Pt(9.5)
    r_h0.font.color.rgb = RGBColor(255, 255, 255)
    r_h1.font.name = 'Arial'
    r_h1.font.bold = True
    r_h1.font.size = Pt(9.5)
    r_h1.font.color.rgb = RGBColor(255, 255, 255)

    set_cell_background(hdr_cells[0], "003366")
    set_cell_background(hdr_cells[1], "003366")
    set_cell_margins(hdr_cells[0], top=30, bottom=30, left=80, right=80)
    set_cell_margins(hdr_cells[1], top=30, bottom=30, left=80, right=80)

    abbrevs = [
        ("AI", "Artificial Intelligence"),
        ("AUC-ROC", "Area Under the Receiver Operating Characteristic Curve"),
        ("CNS", "Central Nervous System"),
        ("CNN", "Convolutional Neural Network"),
        ("CV", "Computer Vision"),
        ("CVAT", "Computer Vision Annotation Tool"),
        ("DINOv2", "Self-Supervised Vision Transformer for Representation Learning"),
        ("GATv2", "Graph Attention Network (Version 2)"),
        ("GNN", "Graph Neural Network"),
        ("IHC", "Immunohistochemistry"),
        ("IoU", "Intersection over Union"),
        ("MAE", "Masked Autoencoder"),
        ("mAP", "Mean Average Precision"),
        ("ML", "Machine Learning"),
        ("MongoDB", "Document-Oriented Database for JSON Metadata & Label Indexing"),
        ("MPNN", "Message Passing Neural Network"),
        ("PBM", "Photobiomodulation Therapy"),
        ("ROS", "Reactive Oxygen Species"),
        ("SAM", "Segment Anything Model"),
        ("SOTA", "State-of-the-Art"),
        ("SSL", "Self-Supervised Learning"),
        ("TBI", "Traumatic Brain Injury"),
        ("U-Net", "Convolutional Network Architecture for Biological Segmentation"),
        ("ViT", "Vision Transformer"),
        ("YOLO", "You Only Look Once (Object Detection Framework)")
    ]

    for abbr, desc in abbrevs:
        row_cells = abbrev_table.add_row().cells
        p0 = row_cells[0].paragraphs[0]
        p1 = row_cells[1].paragraphs[0]
        
        p0.paragraph_format.line_spacing = 1.0
        p0.paragraph_format.space_before = Pt(1)
        p0.paragraph_format.space_after = Pt(1)
        
        p1.paragraph_format.line_spacing = 1.0
        p1.paragraph_format.space_before = Pt(1)
        p1.paragraph_format.space_after = Pt(1)

        r0 = p0.add_run(abbr)
        r0.font.name = 'Arial'
        r0.font.bold = True
        r0.font.size = Pt(9.5)
        r0.font.color.rgb = DARK

        r1 = p1.add_run(desc)
        r1.font.name = 'Arial'
        r1.font.size = Pt(9.5)
        r1.font.color.rgb = DARK

        set_cell_margins(row_cells[0], top=20, bottom=20, left=80, right=80)
        set_cell_margins(row_cells[1], top=20, bottom=20, left=80, right=80)

    doc.add_page_break()

    # 1. MOTIVATION
    add_h1("1. Motivation")
    add_body("Traumatic Brain Injury (TBI) and secondary neurodegenerative disorders represent a major global health crisis, affecting an estimated 69 million individuals annually and standing as a leading cause of long-term neurocognitive impairment (Dewan et al., 2018; Maas et al., 2017). Following primary neurotrauma, a progressive secondary injury cascade ensues, driven by chronic neuroinflammation, oxidative stress, and blood-brain barrier breakdown. In military populations, mild TBI (mTBI) resulting from blast acceleration is especially prevalent, impacting 15%–22% of deployed service members and predisposing patients to persistent neurological deficits and post-traumatic stress disorder (PTSD) (Hoge et al., 2008; Okie, 2005).")
    add_body("Microglia—the primary resident immune sentinels of the Central Nervous System (CNS)—are the central cellular orchestrators of this neuroinflammatory cascade (Salter & Beggs, 2014; Wolf et al., 2017). Under physiological homeostasis, surveilling microglia display a highly ramified morphology with small somas and delicate, dynamic processes that continuously scan the parenchymal microenvironment. Upon encountering mechanical or biochemical stress, microglia undergo rapid morphological metamorphosis, retracting process arbors, enlarging soma volumes, and transitioning into ameboid, hypertrophic, or dystrophic states. Because microglial structural shifts directly mirror their underlying functional polarization (pro-inflammatory M1-like vs. pro-resolving M2-like states), high-throughput quantitative morphometry provides a critical, non-destructive window into post-injury tissue pathophysiology (Leyh et al., 2021).")
    
    # EMBED FIGURE 1: Original Tissue Image
    fig1_file = '/Users/dpeleg/local/MicroGlia/scratch/figures/figure1_original_tissue.jpg'
    add_image_figure(fig1_file, "Figure 1: Example of an original lab microscopy brain tissue slice image (e.g. JPG_VID2724_B1_3_00d07h00m.jpg) containing hundreds of microglial cells marked with cyan contours.", width_inches=5.8)

    add_body("In translational pharmacology and neuro-therapeutics, evaluating the impact of candidate therapeutic drugs and physical modalities—such as Photobiomodulation (PBM) light therapy—on microglial activation and morphological recovery is paramount (Gavish & Houreld, 2019; Hamblin, 2018). In ongoing collaborative research directed by Dr. Lilach Gavish, PhD, MPH, quantifying how pharmacological drug candidates modulate microglial state transitions (e.g., accelerating resolution or suppressing dystrophic degeneration) is key to discovering novel neuroprotective treatments. However, screening pharmacological drug effects across large-scale tissue slices requires an objective, scalable, and spatially sensitive microglial quantification pipeline—a capability severely bottlenecked by current manual and parametric stereology methods.")
    add_body("Recent baseline research at our institution by Presaizen (2026) established a 4-class soma-centric classification pipeline using YOLOv11 and DINOv2 on rat brain slices. While providing an initial proof-of-concept, Presaizen (2026) highlighted fundamental computational limitations: bounding-box object detectors (64x64 crops) restrict analysis to the central soma, discarding up to 70% of distal process arborization. Consequently, bounding-box models suffer from biological misclassification between Resting and Resolution states and fail to detect shattered dystrophic microglia lacking a central soma anchor. Addressing these limitations through a topological, foundation-model-guided AI framework is essential for establishing unbiased morphometric biomarkers to evaluate drug and PBM therapeutic efficacy in translational neurobiology.")

    # 2. RESEARCH QUESTION AND HYPOTHESIS
    add_h1("2. Research Question and Hypothesis")
    add_h2("2.1 Research Questions")
    add_bullet(" Can foundation-model segmentation (e.g., Cellpose 3.0 / SAM-Microscopy) combined with Graph Neural Networks (GNNs) overcome soma-centric bounding box limitations to accurately capture fragmented microglial arborization?", "RQ1:")
    add_bullet(" Does incorporating spatial process topology resolve the persistent biological confusion between Resting and Resolution microglial states?", "RQ2:")
    add_bullet(" Can a continuous, multi-parametric activation index derived from graph topological representations provide superior sensitivity in quantifying pharmacological drug treatments and Photobiomodulation (PBM) therapeutic responses compared to standard discrete classification?", "RQ3:")

    add_h2("2.2 Research Hypothesis")
    add_body("It is hypothesized that transitioning from soma-centric bounding boxes to foundation-model-based polygonal segmentation (Cellpose 3.0 / SAM-Microscopy) and modeling the spatial neighborhood of fragmented distal processes using Graph Neural Networks (GNNs) will resolve the confusion between Resting and Resolution states by capturing the full cellular silhouette. Furthermore, by representing process fragments as nodes in a spatial proximity graph, the framework will significantly increase the recall of dystrophic/shattered cells lacking a unified soma anchor, yielding a sensitive, continuous activation index for evaluating drug and PBM therapeutic efficacy.")

    # 3. RESEARCH OBJECTIVES
    add_h1("3. Research Objectives")
    add_bullet(" Re-annotate the institutional benchmark dataset of 4,874 cells using fine polygonal masks to capture distal processes, beaded arborization, and shattered process fragments excluded by YOLO bounding boxes.", "Objective 1 (Dataset Re-annotation):")
    add_bullet(" Deploy and fine-tune foundation segmentation models (Cellpose 3.0 and SAM-Microscopy) to extract full microglial silhouettes without relying on a central soma anchor.", "Objective 2 (Foundation Segmentation):")
    add_bullet(" Construct spatial proximity graphs connecting soma nodes and process fragment nodes, training a Graph Neural Network (GNN) to reconstruct shattered dystrophic cells into unified biological entities.", "Objective 3 (Graph Topology Construction):")
    add_bullet(" Implement a contrastive self-supervised representation learning space (DINOv2 / Masked Autoencoders) fine-tuned on segmented microglial masks to separate morphologically subtle activation states.", "Objective 4 (Self-Supervised Feature Space):")
    add_bullet(" Formulate a continuous Multi-Parametric Activation Index (0–1 scale) and validate its sensitivity against experimental pharmacological drug-treated, PBM-irradiated, and LPS-challenged rodent brain slices in collaboration with Dr. Lilach Gavish.", "Objective 5 (Pharmacological & Experimental Validation):")

    # 4. LITERATURE REVIEW
    add_h1("4. Literature Review")
    add_h2("4.1 Microglia and Neuroinflammation")
    add_body("Microglia constitute 10%–15% of all glial cells in the CNS and serve as the frontline defense against traumatic, ischemic, and neurodegenerative insults. Under homeostatic conditions, surveilling microglia display an arborized morphology with delicate processes extending tens of micrometers from a small soma. Following TBI, inflammatory signaling cascades induce rapid structural metamorphosis: retracting processes, swelling somas, and transitioning into activated ameboid macrophages capable of phagocytosis.")

    add_h2("4.2 Pharmacological & Photobiomodulation (PBM) Therapeutics")
    add_body("Quantifying microglial morphological response to candidate anti-inflammatory drugs and Photobiomodulation (PBM) therapy is central to translational neuro-therapeutics (Gavish & Houreld, 2019; Hamblin, 2018). In rodent injury models, therapeutic drug compounds and PBM light irradiation stimulate mitochondrial bioenergetics, attenuate reactive oxygen species (ROS), and accelerate microglial polarization from pro-inflammatory (M1-like) to pro-resolving (M2-like) phenotypes. High-throughput quantitative morphometry allows drug discovery researchers to measure dose-dependent morphological shifts across entire brain slices.")

    add_h2("4.3 Morphometric Analysis of Microglia")
    add_body("Traditional microglial morphometry relies on manual thresholding and skeletonization (e.g., ImageJ FracLac or HALO modules). Key quantitative metrics include fractal dimension (D_f), lacunarity, soma-to-cell area ratio, total process length, and number of branch endpoints. While informative, manual and parametric morphometry suffers from severe inter-rater variability, labor intensity, and failure in dense tissue slices with overlapping processes.")

    add_h2("4.4 Automated Detection and Deep Learning")
    add_body("Recent computational advances have applied Convolutional Neural Networks (CNNs) and object detectors (YOLOv8, YOLOv11) to microglial quantification (Anwer et al., 2023; Morera et al., 2024; Hsu et al., 2025 - StainAI). While high-throughput bounding-box detectors excel at counting central somas, they cropped out distal process arborization ($64\\times64$ crops), discarding up to 70% of the morphological information required to distinguish subtle functional states.")

    add_h2("4.5 Foundation Models and Image Segmentation")
    add_body("Foundation models trained on millions of biological images—such as Cellpose 3.0 (Pachitariu & Stringer, 2024) and Segment Anything Model for Microscopy (SAM-Microscopy / SAM 2)—have revolutionized cellular segmentation. By predicting spatial gradient flows and vector fields, Cellpose segment zero-shot cell bodies and extended process branches without bounding-box constraints, providing the ideal input for topological analysis.")

    add_h2("4.6 Graph Neural Networks (GNNs) in Cellular Topology")
    add_body("Graph Neural Networks (GNNs) model complex non-Euclidean spatial relationships. In neurobiology, representing segmented cellular somas and process fragments as nodes in a spatial proximity graph ($G = (V, E)$) enables Message Passing Neural Networks (MPNNs) or Graph Attention Networks (GATs) to learn topological connectivity. GNNs enable the reconstruction of 'shattered' dystrophic microglia—a critical bottleneck in neurodegeneration research.")

    add_h2("4.7 Literature Gap & Summary")
    add_body("Despite rapid progress in deep learning for digital pathology, existing microglial pipelines remain strictly soma-centric and bounding-box constrained. No current framework integrates foundation-model segmentation with spatial GNN topological modeling to resolve Resting vs. Resolution state confusion or reconstruct fragmented dystrophic cells. This project addresses this critical gap.")

    # 5. PRELIMINARY WORK & STUDY RATIONALE
    add_h1("5. Preliminary Work and Study Rationale")
    add_body("This project builds directly upon the baseline M.Sc. thesis project completed at our institution by Tali Presaizen (Jan 2026), supervised by Dr. Sharon Yalov-Handzel and Dr. Lilach Gavish.")
    add_body("The baseline study established a soma-centric annotated dataset of 4,874 microglial cells extracted from phase-contrast and fluorescence microscopy images of PBM-treated rat brain slices. The dataset categorized cells into four discrete morphological states: Resting, Surveilling, Activated, and Resolution.")
    
    # EMBED FIGURE 3: 4 Microglial Cell Activation States Panel
    fig3_file = '/Users/dpeleg/local/MicroGlia/scratch/figures/figure3_cell_states_panel.jpg'
    add_image_figure(fig3_file, "Figure 2: Representative single-cell crop examples across the 4 microglial morphological activation states: (A) Resting (Ramified), (B) Surveilling, (C) Activated (Ameboid), and (D) Resolution.", width_inches=6.0)

    add_body("The baseline architecture utilized a YOLOv11 object detector paired with a DINOv2 Vision Transformer feature extractor for 4-class classification. While achieving respectable baseline metrics (mAP@0.5 ≈ 0.76, macro-F1 ≈ 0.69), the study revealed critical performance bottlenecks:")
    add_bullet(" Soma-centric $64\\times64$ bounding boxes cropped out distal process branches, discarding 70% of morphological information.", "1. Distal Information Loss:")
    add_bullet(" The confusion matrix revealed severe misclassification between Resting and Resolution states (F1 < 0.62), as both states share similar soma sizes but differ vastly in distal process topology.", "2. Resting vs. Resolution Ambiguity:")
    add_bullet(" Dystrophic/shattered microglia in injured brain tissue lack a central soma anchor, causing YOLO to miss up to 45% of fragmented cellular entities.", "3. Failure on Dystrophic Microglia:")
    
    add_body("These baseline findings provide the direct rationale for our proposed Topological AI Pipeline.")

    # 6. PROPOSED METHODOLOGY & TWO-THEME ARCHITECTURE
    add_h1("6. Proposed Methodology & Project Architecture")
    add_body("The project architecture is structured into two core operational themes, ensuring clean separation between data engineering / self-supervised representation learning and downstream model training / pharmacological evaluation:")

    add_h2("THEME 1: Data Preparation, Cleaning, Aggregation, SSL, Storage & Labeling")
    add_bullet(" Automated Whole-Slide Cell Extraction & Cleaning (CLAHE + Scharr/Canny edge fusion + contained sub-cell IoU deduplication).", "Stage 1.1:")
    
    # EMBED FIGURE 2: Whole-Slide Cell Contour Extraction Map
    fig2_file = '/Users/dpeleg/local/MicroGlia/scratch/figures/figure2_extraction_map.jpg'
    add_image_figure(fig2_file, "Figure 3: Example of a whole-slide cell contour extraction grid map (VID2724_A3_4_00d07h00m), illustrating automated single-cell extraction and side-by-side pipeline evaluation.", width_inches=5.8)

    add_bullet(" Hybrid Storage Architecture (MongoDB JSON document store for metadata, spatial BBoxes [x, y, w, h], cluster IDs, and active labels + HDF5 binary containers sharded by original Image ID).", "Stage 1.2:")
    add_bullet(" Lab Stain Normalization Engine (Macenko optical density matrix factorization BEFORE SSL pre-training to eliminate IHC color shifts).", "Stage 1.3:")
    add_bullet(" In-Domain Self-Supervised Pre-Training (DINOv2 self-distillation + MAE patch reconstruction on 1M+ stain-normalized crops).", "Stage 1.4:")
    add_bullet(" Unsupervised Feature Space Pre-Clustering (UMAP + HDBSCAN partitioning embeddings into ~100 morphometric clusters).", "Stage 1.5:")
    add_bullet(" Active Bulk Labeling & Uncertainty Sampling (MongoDB-indexed 1-click bulk cluster verification + top 5% entropy expert sampling).", "Stage 1.6:")

    add_h2("THEME 2: Training, Classification, Counting & Evaluation")
    add_bullet(" Deterministic Spatial Graph Construction (Connecting Soma Nodes V_soma and Fragment Nodes V_fragment via Delaunay/k-NN edges using spatial BBox coordinates).", "Stage 2.1:")
    add_bullet(" Multi-Task Joint Model Training (Pre-trained SSL DINOv2 ViT backbone + GATv2 Graph Encoder trained with Focal + Contrastive + Reconstruction loss).", "Stage 2.2:")
    add_bullet(" Whole-Slide High-Throughput Inference & Seam NMS Deduplication (Parallel overlapping tile inference + border NMS).", "Stage 2.3:")
    add_bullet(" Per-State Cell Counting & Continuous Activation Index Computation (Computing discrete 5-state counts and 0.00–1.00 continuous Activation Score).", "Stage 2.4:")
    add_bullet(" Pharmacological Drug & PBM Response Analytics Platform (Dose-response sensitivity correlation r/rho for Dr. Lilach Gavish's screening pipeline).", "Stage 2.5:")

    # 7. WORK PLAN & STAGES (MATCHING LINOY'S FORMAT)
    add_h1("7. Work Plan and Project Stages")
    add_body("The proposed research will be executed across six structured project stages over a total estimated duration of 26 weeks (~6.5 months), adhering to the departmental proposal guidelines:")

    stages_data = [
        ("Stage 1 – Theme 1: Dataset Extraction, MongoDB Setup & Stain Normalization",
         "Run automated cell extraction (extract_cells.py). Set up MongoDB document store for JSON metadata, spatial BBoxes ([x, y, w, h]), and active labels. Perform Macenko stain normalization on all 1,000,000+ cell crops sharded by original Image ID (IMAGE_ID_cells.h5).",
         "Establish a clean, stain-normalized single-cell crop repository indexed in MongoDB.",
         "4 weeks"),
        ("Stage 2 – Theme 1: SSL Pre-Training & Active Cluster Labeling",
         "Pre-train DINOv2 and MAE backbones on 1M+ stain-normalized crops. Run HDBSCAN clustering and annotate 10,000–50,000 cells using MongoDB-driven active bulk cluster verification in CVAT.",
         "Deploy a domain-specific SSL feature encoder and establish a gold-standard labeled benchmark dataset.",
         "4 weeks"),
        ("Stage 3 – Theme 2: Spatial GNN Graph Construction & Topology",
         "Construct physical spatial proximity graphs (G=(V,E)) connecting soma nodes and process fragment nodes using BBox spatial coordinates ([x, y, w, h]). Implement and train GATv2/MPNN GNN architectures.",
         "Reconstruct shattered dystrophic microglia into single biological entities and capture full arborization topology.",
         "5 weeks"),
        ("Stage 4 – Theme 2: Multi-Task Joint Model Fine-Tuning",
         "Fine-tune the joint DINOv2 ViT + GATv2 GNN architecture on labeled data using combined Focal, Contrastive, and Dystrophic Reconstruction losses.",
         "Achieve state-of-the-art per-class classification accuracy (Macro-F1 > 0.94) and resolve Resting vs. Resolution state ambiguity.",
         "4 weeks"),
        ("Stage 5 – Theme 2: High-Throughput Whole-Slide Inference Engine",
         "Build the whole-slide inference pipeline with overlapping 1024x1024 tile processing, BBox-driven Non-Maximum Suppression (NMS), per-state counting, and continuous Activation Index (0.00–1.00) computation.",
         "Deliver a fast, automated whole-slide cell counting engine.",
         "5 weeks"),
        ("Stage 6 – Theme 2: Pharmacological Validation, Thesis Writing & Defense",
         "Perform statistical sensitivity validation (Pearson r, Spearman rho) on drug-treated and PBM-irradiated rat brain slices with Dr. Lilach Gavish. Write final M.Sc. thesis and defend before academic committee.",
         "Complete and submit the final M.Sc. thesis document and defend the research.",
         "4 weeks")
    ]

    for stage_title, task, goal, duration in stages_data:
        add_h2(stage_title)
        add_bullet(task, "Task: ")
        add_bullet(goal, "Goal: ")
        add_bullet(duration, "Estimated Duration: ")

    # 8. EVALUATION PLAN
    add_h1("8. Evaluation Plan and Benchmarking")
    add_body("The framework will be evaluated across three complementary quantitative tiers:")
    add_bullet(" Dice Coefficient, Intersection over Union (IoU), and Boundary-F1 score compared against manual polygonal ground truth.", "1. Segmentation Metrics:")
    add_bullet(" Macro-F1 score, Per-class Precision/Recall, and Confusion Matrix analysis across Resting, Surveilling, Activated, Resolution, and Dystrophic states (benchmarked against YOLOv11 baseline F1=0.69).", "2. Morphological Classification:")
    add_bullet(" Pearson correlation ($r$) and Spearman rank ($\\rho$) between the computed Activation Index (0–1) and biological pharmacological drug dosage / PBM light fluence ($J/cm^2$) across tissue slices.", "3. Biological & Clinical Sensitivity:")

    # 9. EXPECTED CONTRIBUTION
    add_h1("9. Expected Scientific and Technological Contribution")
    add_bullet(" First AI framework combining foundation-model segmentation with Graph Neural Networks for microglial morphometry.", "1. Novel Methodological Paradigm:")
    add_bullet(" Solves the critical bottleneck of detecting shattered dystrophic microglia lacking a central soma anchor.", "2. Dystrophic Microglia Reconstruction:")
    add_bullet(" Provides an open-source, fragment-first polygonal dataset of 4,874 microglial cells for the research community.", "3. Open-Source Benchmark Dataset:")
    add_bullet(" Delivers a scalable, reproducible tool for quantifying pharmacological drug impact and Photobiomodulation (PBM) neuroprotective efficacy in translational neurobiology under the direction of Dr. Lilach Gavish.", "4. Translational Pharmacology Impact:")

    # 10. REFERENCES
    add_h1("10. References")
    references_list = [
        "Anwer, D. M., Gubinelli, F., Kurt, Y. A., et al. (2023). A comparison of machine learning approaches for the quantification of microglial cells in the brain of mice, rats and non-human primates. PLOS ONE, 18(4), e0284480.",
        "Dewan, M. C., Rattani, A., Gupta, S., et al. (2018). Estimating the global incidence of traumatic brain injury. Journal of Neurosurgery, 130(4), 1080-1097.",
        "Gavish, L., & Houreld, N. N. (2019). Therapeutic Efficacy of Photobiomodulation (PBM) in Wound Healing and Neuroinflammation. Photomedicine and Laser Surgery, 37(3), 150-162.",
        "Hamblin, M. R. (2018). Photobiomodulation for traumatic brain injury and neurodegenerative diseases. Photonics & Lasers in Medicine, 7(3), 231-244.",
        "He, K., Chen, X., Xie, S., et al. (2022). Masked autoencoders are scalable vision learners. IEEE/CVF CVPR, 16000-16009.",
        "Hoge, C. W., McGurk, D., Thomas, J. L., et al. (2008). Mild traumatic brain injury in U.S. Soldiers returning from Iraq. New England Journal of Medicine, 358(5), 453-463.",
        "Hsu, C.-H., Hsu, Y.-Y., Chang, B.-M., et al. (2025). StainAI: quantitative mapping of stained microglia and insights into brain-wide neuroinflammation and therapeutic effects in cardiac arrest. Communications Biology, 8, 7926.",
        "Kim, J., Pavlidis, P., & Vogel Ciernia, A. (2024). Development of a High-Throughput Pipeline to Characterize Microglia Morphological States at a Single-Cell Resolution. eNeuro, 11(6), ENEURO.0010-24.2024.",
        "Kirillov, A., Mintun, E., Ravi, N., et al. (2023). Segment Anything. Proceedings of the IEEE/CVF International Conference on Computer Vision (ICCV), 4015-4026.",
        "Leyh, J., Schafer, M. K., et al. (2021). Microglial morphodynamics in traumatic brain injury and recovery. Glia, 69(8), 1950-1965.",
        "Maas, A. I., Menon, D. K., Adelson, P. D., et al. (2017). Traumatic brain injury: integrated approaches to improve prevention, clinical care, and research. The Lancet Neurology, 16(12), 987-1048.",
        "Macenko, M., Niethammer, M., Marron, J. S., et al. (2009). A method for normalizing histology slides for quantitative analysis. IEEE ISBI, 1107-1110.",
        "Morera, H., Dave, P., Kolinko, Y., et al. (2024). A novel deep learning-based method for automatic stereology of microglia cells from low magnification images. Neurotoxicology and Teratology, 102, 107336.",
        "Oquab, M., Darcet, T., Moutakanni, T., et al. (2023). DINOv2: Learning Robust Visual Features Without Supervision. arXiv preprint arXiv:2304.07193.",
        "Pachitariu, M., & Stringer, C. (2024). Cellpose 3.0: accurate segmentation of biological images using foundation models. Nature Methods, 21(4), 701-710.",
        "Presaizen, T. (2026). AI-Powered Microglial Classification for Activation Scoring. Master's Thesis, School of Data Science: Intelligent Systems, Afeka Academic College of Engineering & Hebrew University of Jerusalem.",
        "Salter, M. W., & Beggs, S. (2014). Sublime microglia: expanding roles for the guardians of the CNS. Cell, 158(1), 15-24.",
        "Velickovic, P., Cucurull, G., Casanova, A., et al. (2018). Graph Attention Networks. International Conference on Learning Representations (ICLR).",
        "Wolf, S. A., Boddeke, H. W., & Kettenmann, H. (2017). Microglia in Physiology and Pathology. Physiological Reviews, 97(4), 1339-1393.",
        "Xiong, H., Zheng, S., Qi, X., Liu, J. (2025). μGlia-Flow, an automatic workflow for microglia segmentation and classification. Journal of Neuroscience Methods, 402, 110022.",
        "Zähringer, A., Vinnakota, J. M., Wertheimer, T., et al. (2025). AIstain: Enhancing microglial phagocytosis analysis through deep learning. Cell Reports Methods, 5(11), 101207."
    ]

    for ref in references_list:
        p_ref = doc.add_paragraph()
        p_ref.paragraph_format.left_indent = Inches(0.5)
        p_ref.paragraph_format.first_line_indent = Inches(-0.5)
        p_ref.paragraph_format.space_after = Pt(4)
        run = p_ref.add_run(ref)
        run.font.name = 'Arial'
        run.font.size = Pt(10)
        run.font.color.rgb = DARK

    # Save DOCX files
    out_local = '/Users/dpeleg/local/MicroGlia/research-proposal-final.docx'
    out_downloads = '/Users/dpeleg/Downloads/research-proposal-final.docx'
    doc.save(out_local)
    doc.save(out_downloads)
    print(f"Updated Figure-Embedded DOCX Proposal saved to:\n  - {out_local}\n  - {out_downloads}")

if __name__ == "__main__":
    create_research_proposal()
